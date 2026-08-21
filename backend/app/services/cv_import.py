"""CV import — turn an uploaded CV into master-profile suggestions.

Two paths, mirroring services/normalizer.py (spec §20, §21):
  1. LLM path (preferred): strict extract-only prompt → JSON → validated.
  2. No-LLM path (fallback): deterministic regex/heuristics, so the feature
     keeps working with zero API keys configured.

Nothing here writes to the master profile. The caller returns these as
*suggestions*; the user confirms them through the normal PUT /profile form.
That preserves the rule on DocumentExtraction: never auto-promote to the
master profile until the user confirms.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from pydantic import BaseModel

from app.core.llm import LLMError, get_provider
from app.schemas.profile import (
    CvImportResult,
    CvSuggestedProfile,
    EducationIn,
    ExperienceIn,
    SkillIn,
)
from app.services.prompts import build_message

logger = logging.getLogger("careerpilot.cv_import")

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# Cap the text sent to the LLM / regex engine. A CV is a few pages; anything
# beyond this is almost certainly a mis-uploaded document.
MAX_CV_CHARS = 40_000


class CvImportError(RuntimeError):
    """Raised when the CV cannot be read at all."""


# ── Text extraction ─────────────────────────────────────────────


def extract_text(path: Path, suffix: str | None = None) -> str:
    """Extract plain text from a PDF, DOCX or plain-text CV."""
    suffix = (suffix or path.suffix).lower()

    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise CvImportError(f"Unsupported CV format {suffix or '(none)'}")

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of blank lines but keep single newlines (layout matters).
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", "  ", text)
    return text.strip()[:MAX_CV_CHARS]


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise CvImportError(
            "PDF support requires the 'pypdf' package. Install requirements.txt."
        ) from exc
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise CvImportError(
                    "This PDF is password-protected. Remove the password and retry."
                ) from exc
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except CvImportError:
        raise
    except Exception as exc:
        logger.warning("PDF text extraction failed: %s", type(exc).__name__)
        raise CvImportError("Could not read this PDF. Try DOCX, or a text-based PDF.") from exc


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise CvImportError("DOCX support requires the 'python-docx' package.") from exc
    try:
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("DOCX text extraction failed: %s", type(exc).__name__)
        raise CvImportError("Could not read this DOCX file.") from exc


# ── Parsing ─────────────────────────────────────────────────────


def parse_cv(text: str) -> CvImportResult:
    """Parse CV text into profile suggestions. LLM first, deterministic fallback."""
    if not text.strip():
        return CvImportResult(
            warnings=[
                "No text could be read from this file. If it is a scanned CV, "
                "the text must be selectable — try a DOCX or a text-based PDF."
            ]
        )

    try:
        result = _parse_with_llm(text)
        result.parser = "llm"
    except (LLMError, ValueError, FileNotFoundError) as exc:
        logger.info("CV LLM parse unavailable (%s); using deterministic parser", type(exc).__name__)
        result = _parse_deterministic(text)
        result.parser = "fallback"
    except Exception:  # pragma: no cover - never let a parser bug 500 the upload
        logger.exception("Unexpected CV LLM parse failure; using deterministic parser")
        result = _parse_deterministic(text)
        result.parser = "fallback"

    result.filled_fields = [
        name
        for name, value in result.profile.model_dump().items()
        if isinstance(value, str) and value.strip()
    ]
    if not result.filled_fields and not (result.education or result.experience or result.skills):
        result.warnings.append(
            "We could not confidently read any details from this CV. "
            "Please fill the form in manually."
        )
    return result


def _parse_with_llm(text: str) -> CvImportResult:
    system, user = build_message("cv_extraction", cv_text=text)
    parsed = get_provider().complete_json(system, user, _LlmCvPayload, task="cv_extraction")
    return CvImportResult(
        profile=parsed.profile or CvSuggestedProfile(),
        education=parsed.education,
        experience=parsed.experience,
        skills=_dedupe_skills(parsed.skills),
    )


# ── Deterministic fallback ──────────────────────────────────────

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
# Kenyan and international formats: 0712345678, +254712345678, (020) 123 4567
_PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d[\d\s().-]{7,17}\d)(?!\w)")

_SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "summary": ("professional summary", "personal statement", "profile summary",
                "career objective", "objective", "about me", "summary", "profile"),
    "education": ("education and qualifications", "academic qualifications",
                  "educational background", "qualifications", "education"),
    "experience": ("work experience", "professional experience", "teaching experience",
                   "employment history", "career history", "experience", "employment"),
    "skills": ("key skills", "core competencies", "technical skills",
               "skills and competencies", "competencies", "skills"),
    "certifications": ("certifications", "certificates", "licences", "licenses",
                       "professional registration", "memberships"),
}

_DEGREE_RE = re.compile(
    r"\b(bachelor|master|doctor|phd|d\.?phil|b\.?ed|b\.?sc|b\.?a\b|m\.?ed|m\.?sc|m\.?a\b|"
    r"mba|diploma|certificate|higher national|hnd|kcse|kcpe)\b",
    re.I,
)
_CLASSIFICATION_RE = re.compile(
    r"\b(first class(?: honou?rs)?|second class(?: upper| lower)?(?: division)?|"
    r"upper second|lower second|third class|distinction|credit|merit|pass|"
    r"gpa\s*[:\-]?\s*\d(?:\.\d+)?|mean grade\s*[:\-]?\s*[A-E][+-]?)\b",
    re.I,
)
_YEAR_RANGE_RE = re.compile(
    r"((?:19|20)\d{2})\s*(?:[-–—]|to)\s*((?:19|20)\d{2}|present|current|date|to date)",
    re.I,
)
_CURRENT_RE = re.compile(r"\b(present|current|to date|ongoing)\b", re.I)
_REGISTRATION_RE = re.compile(
    r"\b(tsc|teachers? service commission|registration (?:no|number)|licence no|license no|"
    r"reg(?:istration)?\.? ?no)\b[^\n]*",
    re.I,
)
_ROLE_AT_ORG_RE = re.compile(
    r"^(?P<role>[A-Z][\w /&.'-]{2,60}?)\s*(?:,|—|–|-|\bat\b|\|)\s*(?P<org>[A-Z][\w /&.,'-]{2,80})",
)


def _split_sections(text: str) -> dict[str, list[str]]:
    """Group CV lines under normalised section headings."""
    lines = [ln.rstrip() for ln in text.split("\n")]
    sections: dict[str, list[str]] = {"_header": []}
    current = "_header"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        heading = _match_heading(stripped)
        if heading:
            current = heading
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(stripped)
    return sections


# Keywords used when a heading is not an exact alias. CVs vary wildly
# ("PROFESSIONAL SKILLS", "CORE SKILLS", "AREAS OF EXPERTISE"...), so an exact
# list will always miss some. Order matters: the first match wins.
_SECTION_KEYWORDS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("education", ("education", "academic", "qualification", "schooling")),
    ("experience", ("experience", "employment", "work history", "career history",
                    "professional background", "teaching practice")),
    ("skills", ("skill", "competenc", "expertise", "proficienc", "strengths")),
    ("certifications", ("certification", "certificate", "licence", "license",
                        "registration", "membership", "accreditation")),
    ("summary", ("summary", "profile", "objective", "personal statement", "about me")),
)


def _looks_like_heading(line: str) -> bool:
    """Heading-ish formatting: short, and capitalised or colon-terminated."""
    stripped = line.strip()
    if not (2 < len(stripped) <= 45):
        return False
    if stripped.endswith(":"):
        return True
    letters = [c for c in stripped if c.isalpha()]
    if letters and all(c.isupper() for c in letters):
        return True          # ALL CAPS
    words = stripped.split()
    # Title Case and few words, e.g. "Key Skills"
    return len(words) <= 4 and all(w[:1].isupper() for w in words if w[:1].isalpha())


def _match_heading(line: str) -> str | None:
    """Return the canonical section name if this line looks like a heading."""
    if len(line) > 60:
        return None
    cleaned = re.sub(r"[^a-z& ]", "", line.lower()).strip()
    if not cleaned:
        return None
    # 1. exact alias (fast path, unambiguous)
    for canonical, aliases in _SECTION_ALIASES.items():
        if cleaned in aliases:
            return canonical
    # 2. keyword match, but only for lines that are formatted like a heading,
    #    so a sentence such as "Skills gained during my placement" is not one.
    if not _looks_like_heading(line):
        return None
    for canonical, keywords in _SECTION_KEYWORDS:
        if any(word in cleaned for word in keywords):
            return canonical
    return None


def _parse_deterministic(text: str) -> CvImportResult:
    sections = _split_sections(text)
    header = sections.get("_header", [])
    profile = CvSuggestedProfile()

    email = _EMAIL_RE.search(text)
    if email:
        profile.email = email.group(0)

    phone = _find_phone(text)
    if phone:
        profile.phone = phone

    profile.full_name = _guess_name(header)
    profile.profession = _guess_profession(header, profile.full_name)
    profile.location = _guess_location(header)

    nationality = re.search(r"\bnationality\s*[:\-]\s*([A-Za-z ]{3,30})", text, re.I)
    if nationality:
        profile.nationality = nationality.group(1).strip()

    summary_lines = sections.get("summary", [])
    if summary_lines:
        profile.summary = " ".join(summary_lines)[:1500].strip()

    profile.professional_registration = _find_registration(text)

    return CvImportResult(
        profile=profile,
        education=_parse_education(sections.get("education", [])),
        experience=_parse_experience(sections.get("experience", [])),
        skills=_dedupe_skills(_parse_skills(sections.get("skills", []))),
    )


def _find_phone(text: str) -> str | None:
    for match in _PHONE_RE.finditer(text):
        candidate = match.group(0).strip()
        digits = re.sub(r"\D", "", candidate)
        # Reject year ranges and other numeric noise.
        if 9 <= len(digits) <= 15 and not re.fullmatch(r"(?:19|20)\d{2}\D*(?:19|20)\d{2}", candidate):
            return candidate
    return None


def _guess_name(header: list[str]) -> str | None:
    for line in header[:6]:
        if _EMAIL_RE.search(line) or _PHONE_RE.search(line):
            continue
        if any(ch.isdigit() for ch in line):
            continue
        words = line.split()
        if not (2 <= len(words) <= 5):
            continue
        if len(line) > 60:
            continue
        # "JOHN GICHAGA" or "John Gichaga", not "Curriculum Vitae"
        if line.lower() in {"curriculum vitae", "resume", "cv", "personal details"}:
            continue
        if all(w[:1].isupper() for w in words if w[:1].isalpha()):
            return " ".join(w.capitalize() if w.isupper() else w for w in words)
    return None


def _guess_profession(header: list[str], name: str | None) -> str | None:
    for line in header[:8]:
        if name and line.lower() == name.lower():
            continue
        if _EMAIL_RE.search(line) or _PHONE_RE.search(line):
            continue
        if 3 < len(line) <= 80 and re.search(
            r"\b(teacher|tutor|lecturer|engineer|developer|analyst|manager|officer|"
            r"consultant|specialist|nurse|accountant|designer|scientist|administrator|"
            r"coordinator|instructor|trainer)\b",
            line,
            re.I,
        ):
            return line.strip(" -–—|")
    return None


def _guess_location(header: list[str]) -> str | None:
    for line in header[:8]:
        match = re.search(r"\b([A-Z][a-z]+(?: [A-Z][a-z]+)*),\s*(Kenya|Uganda|Tanzania|[A-Z][a-z]+)\b", line)
        if match and not _EMAIL_RE.search(line):
            return match.group(0)
    return None


def _parse_education(lines: list[str]) -> list[EducationIn]:
    entries: list[EducationIn] = []
    for line in lines:
        if not _DEGREE_RE.search(line):
            continue
        entry = EducationIn(degree=_degree_fragment(line))
        institution = re.search(
            r"(?:\bat\b|,|-|–|—|\|)\s*"
            r"([A-Z][\w .'&-]*?(?:University|College|School|Institute|Polytechnic)[\w .'&-]*)",
            line,
        )
        if institution:
            entry.institution = _clean_fragment(
                _CLASSIFICATION_RE.sub("", _strip_dates(institution.group(1)))
            )
        classification = _CLASSIFICATION_RE.search(line)
        if classification:
            entry.classification = classification.group(0).strip()
        years = _YEAR_RANGE_RE.search(line)
        if years:
            entry.start_date = years.group(1)
            end = years.group(2)
            if _CURRENT_RE.fullmatch(end) or end.lower() in {"present", "current", "date", "to date"}:
                entry.is_current = True
            else:
                entry.end_date = end
        elif (single := re.search(r"\b((?:19|20)\d{2})\b", line)):
            entry.end_date = single.group(1)
        entries.append(entry)
    return entries[:15]


def _parse_experience(lines: list[str]) -> list[ExperienceIn]:
    entries: list[ExperienceIn] = []
    for line in lines:
        years = _YEAR_RANGE_RE.search(line)
        role_org = _ROLE_AT_ORG_RE.match(line)
        if not (years or role_org):
            continue
        entry = ExperienceIn()
        if role_org:
            entry.role = _strip_dates(role_org.group("role"))
            entry.organization = _strip_dates(role_org.group("org"))
        else:
            entry.role = _strip_dates(line)
        if years:
            entry.start_date = years.group(1)
            end = years.group(2)
            if end.lower() in {"present", "current", "date", "to date"}:
                entry.is_current = True
            else:
                entry.end_date = end
        entries.append(entry)
    return entries[:20]


def _parse_skills(lines: list[str]) -> list[SkillIn]:
    skills: list[SkillIn] = []
    for line in lines:
        # Skills are usually comma/bullet/pipe separated on one or more lines.
        for part in re.split(r"[,•·|;/]|\s{2,}", line):
            name = part.strip(" -–—\t")
            if 2 <= len(name) <= 60 and not name.endswith(":") and re.search(r"[A-Za-z]", name):
                skills.append(SkillIn(name=name))
    return skills


def _dedupe_skills(skills: list[SkillIn]) -> list[SkillIn]:
    seen: set[str] = set()
    unique: list[SkillIn] = []
    for skill in skills:
        key = skill.name.strip().lower()
        if key and key not in seen:
            seen.add(key)
            unique.append(skill)
    return unique[:40]


def _find_registration(text: str) -> str | None:
    """Prefer a registration line that carries an actual number/reference."""
    matches = [m.group(0).strip() for m in _REGISTRATION_RE.finditer(text)]
    if not matches:
        return None
    with_number = [m for m in matches if re.search(r"\d", m)]
    return (with_number[0] if with_number else matches[0])[:300]


def _strip_dates(value: str) -> str:
    """Remove trailing year ranges/single years left over from a CV line."""
    value = _YEAR_RANGE_RE.sub("", value)
    value = re.sub(r"\b(?:19|20)\d{2}\b", "", value)
    return _clean_fragment(value)


def _degree_fragment(line: str) -> str:
    """Keep the qualification itself, dropping institution/grade/date noise."""
    candidate = re.split(r"\s*(?:,|\||—|–| - )\s*", line)[0]
    if not _DEGREE_RE.search(candidate):
        candidate = line
    candidate = _CLASSIFICATION_RE.sub("", candidate)
    return _strip_dates(candidate)


def _clean_fragment(value: str) -> str:
    return re.sub(r"\s{2,}", " ", value).strip(" ,;:-–—\t")


class _LlmCvPayload(BaseModel):
    """Internal model describing the LLM JSON contract."""

    profile: CvSuggestedProfile | None = None
    education: list[EducationIn] = []
    experience: list[ExperienceIn] = []
    skills: list[SkillIn] = []
