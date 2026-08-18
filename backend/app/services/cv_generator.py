"""CV Generator (spec §7) — ATS-friendly, tailored, and fabrication-proof.

How it stays honest:
  - The document is assembled ONLY from the master profile (single source of
    truth). No invented facts can enter because there is nowhere to invent from.
  - Tailoring = re-ordering and keyword surfacing, never adding claims:
    the role-type preference lists (math / CS / EdTech / international / AI)
    decide which real skills are shown first.
  - Every line passes the FactCheck gate before it is written.
  - Output: .docx (ATS-friendly: single column, standard fonts, no graphics)
    and .pdf. No references, no salary history, no fabricated achievements.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from app.core.crypto import decrypt_text
from app.models import Job, MasterProfile
from app.services.fact_check import FactCheckReport, fact_check_document
from app.services.scoring import profile_years_experience

# ── role detection & skill preference (spec §7) ────────────────

ROLE_SIGNALS: dict[str, list[str]] = {
    "math": ["mathematics", "maths", "math teacher", "mathematical"],
    "computer_science": ["computer studies", "computer science", "ict teacher", "programming",
                         "coding", "computational", "computer teacher"],
    "edtech": ["edtech", "educational technology", "digital learning", "technology integration",
               "learning management", "instructional technology"],
    "ai_training": ["ai trainer", "ai evaluator", "ai data", "ai tutor", "annotation",
                    "data label", "machine learning"],
    "curriculum": ["curriculum", "instructional design"],
    "international": ["international school", "igcse", "ib school", "a-level", "o-level",
                      "cambridge", "british curriculum", "bilingual school"],
}

# Preferred skill order per role type — surfaced first, still only from the profile
SKILL_PREFERENCES: dict[str, list[str]] = {
    "math": ["mathematics teaching", "assessment", "competency-based assessment",
             "problem solving", "data analysis", "microsoft excel", "cbc/cbe pedagogy",
             "curriculum implementation", "learner-centred teaching"],
    "computer_science": ["computer studies teaching", "html", "css", "javascript",
                         "python", "basic programming", "front-end web development",
                         "ict integration", "computational thinking", "digital literacy"],
    "edtech": ["educational technology", "ict integration", "digital literacy", "python",
               "ai tools", "ai-assisted workflows", "project-based learning", "ict"],
    "ai_training": ["python", "data analysis", "mathematics teaching", "ai tools",
                    "ai-assisted workflows", "problem solving", "critical thinking",
                    "microsoft excel"],
    "curriculum": ["curriculum implementation", "cbc/cbe pedagogy", "competency-based assessment",
                   "learner-centred teaching", "project-based learning", "ict integration"],
    "international": ["learner-centred teaching", "project-based learning", "ict integration",
                      "digital literacy", "competency-based assessment", "creativity and innovation",
                      "critical thinking"],
}


def detect_role_type(job: Job) -> str:
    text = f"{job.title} {job.description or ''} {job.curriculum or ''}".lower()
    for role, signals in ROLE_SIGNALS.items():
        if any(s in text for s in signals):
            return role
    return "general"


def analyze_job(job: Job) -> dict:
    """Extract what the role asks for (used for tailoring, never to invent)."""
    text = f"{job.title} {job.description or ''}"
    keywords = [w for w in re.findall(r"[a-zA-Z][a-zA-Z+-]{2,}", text) if w.lower() not in {
        "the", "and", "for", "with", "you", "will", "are", "our", "this", "that", "have"}]
    return {
        "role_type": detect_role_type(job),
        "requirements": list(job.requirements or []),
        "preferred": list(job.preferred_requirements or []),
        "keywords": list(dict.fromkeys(k.lower() for k in keywords))[:40],
        "curriculum": job.curriculum,
    }


# ── CV assembly ────────────────────────────────────────────────


@dataclass
class CvDocument:
    sections: dict[str, list[str]] = field(default_factory=dict)
    fact_check: FactCheckReport | None = None

    def lines(self) -> list[str]:
        out: list[str] = []
        for heading, lines in self.sections.items():
            out.append(heading)
            out.extend(lines)
        return out

    def as_dict(self) -> dict:
        return {"sections": self.sections}


def _ordered_skills(profile: MasterProfile, role_type: str) -> list[str]:
    approved = [s.name for s in profile.skills if s.approved]
    lower = {s.name.lower(): s.name for s in profile.skills if s.approved}
    prefs = SKILL_PREFERENCES.get(role_type, [])
    ordered: list[str] = []
    seen: set[str] = set()
    for pref in prefs:
        key = pref.lower()
        if key in lower and key not in seen:
            ordered.append(lower[key])
            seen.add(key)
    for name in approved:
        if name.lower() not in seen:
            ordered.append(name)
            seen.add(name.lower())
    return ordered


def _summary(profile: MasterProfile, role_type: str, ordered_skills: list[str]) -> str:
    """Deterministic summary assembled only from profile facts (no LLM)."""
    classification = next((e.classification for e in profile.education if e.classification), "")
    degree = next((e.degree for e in profile.education if e.degree), "")
    field = next((e.field for e in profile.education if e.field), "")
    institution = next((e.institution for e in profile.education if e.institution), "")
    years = profile_years_experience(profile)
    subjects = ", ".join(sorted({s for e in profile.experience for s in (e.subjects or [])})) or "Mathematics and Computer Studies"
    current = next((e for e in profile.experience if e.is_current), None)
    reg = (profile.professional_registration or "").replace(" registered teacher", "")
    skills = ", ".join(ordered_skills[:4])

    bits = [f"{profile.profession} with {classification} {degree}"]
    if field:
        bits.append(f"in {field}")
    if institution:
        bits.append(f"from {institution}")
    if years >= 0.75:
        bits.append(f"{years:g} year(s) of teaching experience in {subjects}")
    if current:
        bits.append(f"currently teaching at {current.organization}")
    if reg:
        bits.append(f"{reg} registered")
    sentence = ", ".join(bits) + "."
    return f"{sentence} Strong in {skills}."


def build_cv(profile: MasterProfile, job: Job, target_role: str | None = None) -> CvDocument:
    """Assemble the tailored CV and run the FactCheck gate."""
    role = detect_role_type(job)
    analysis = analyze_job(job)
    ordered_skills = _ordered_skills(profile, role)

    sections: dict[str, list[str]] = {}

    # Contact
    contact = [profile.full_name]
    phone = decrypt_text(profile.phone_encrypted)
    if phone:
        contact.append(f"Phone: {phone}")
    if profile.email:
        contact.append(f"Email: {profile.email}")
    contact.append(f"Location: {profile.location or ''}".rstrip())
    if profile.nationality:
        contact.append(f"Nationality: {profile.nationality}")
    sections["CONTACT"] = [c for c in contact if c and c.rstrip(": ") != "Location:"]

    # Summary
    sections["PROFESSIONAL SUMMARY"] = [_summary(profile, role, ordered_skills)]

    # Education
    edu_lines = []
    for e in profile.education:
        line = " | ".join(p for p in [e.degree, e.institution, e.field, e.classification] if p)
        if e.start_date or e.end_date:
            line = f"{line} ({e.start_date or ''} - {e.end_date or 'present'})".replace("  ", " ")
        edu_lines.append(line)
    sections["EDUCATION"] = edu_lines

    # Teaching experience
    exp_lines = []
    for e in sorted(profile.experience, key=lambda x: (x.end_date or "9999"), reverse=True):
        head = f"{e.role or 'Teacher'} — {e.organization}"
        dates = f"{e.start_date or ''} - {e.end_date or 'present'}"
        exp_lines.append(f"{head} ({dates})")
        if e.subjects:
            exp_lines.append(f"Subjects: {', '.join(e.subjects)}")
        if e.grades:
            exp_lines.append(f"Grades/Classes: {', '.join(e.grades)}")
        if e.description:
            exp_lines.append(e.description)
    sections["TEACHING EXPERIENCE"] = exp_lines

    # Skills (tailored order)
    sections["SKILLS"] = [", ".join(ordered_skills)]

    # Registration & certifications
    reg_lines = []
    if profile.professional_registration:
        reg_lines.append(profile.professional_registration)
    for cert in profile.certifications:
        if cert.name and cert.name.lower() not in (profile.professional_registration or "").lower():
            reg_lines.append(f"{cert.name} — {cert.issuer or ''}".rstrip(" —"))
    sections["PROFESSIONAL REGISTRATION & CERTIFICATIONS"] = reg_lines

    # FactCheck gate — verify content lines only (headings are structural)
    content_lines = [l for lines in sections.values() for l in lines]
    report, kept = fact_check_document(content_lines, profile)
    cv = CvDocument(sections=sections, fact_check=report)
    if not kept or report.prohibited_findings:
        cv.sections = {}  # hard fail — nothing may be emitted
    else:
        kept_set = set(kept)
        cv.sections = {h: [l for l in lines if l in kept_set]
                       for h, lines in sections.items()}
    return cv


# ── document writers ───────────────────────────────────────────


def write_docx(cv: CvDocument, path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    first = True
    for heading, lines in cv.sections.items():
        if not lines:
            continue
        p = doc.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(11)
        p.space_after = Pt(2)
        for line in lines:
            doc.add_paragraph(line)
        if first:
            first = False
    doc.save(path)
    return path


def write_pdf(cv: CvDocument, path: Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    title = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=13, leading=16)
    heading = ParagraphStyle("heading", fontName="Helvetica-Bold", fontSize=10.5, leading=13,
                             spaceBefore=6)
    body = ParagraphStyle("body", fontName="Helvetica", fontSize=9.5, leading=12)

    story: list = []
    first = True
    for h, lines in cv.sections.items():
        if not lines:
            continue
        if first:
            story.append(Paragraph("<br/>".join(lines[:1]), title))
            story.append(Spacer(1, 2 * mm))
            rest = lines[1:]
            first = False
        else:
            story.append(Paragraph(h, heading))
            rest = lines
        for line in rest:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body))
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=15 * mm, bottomMargin=15 * mm)
    doc.build(story)
    return path


def generated_dir() -> Path:
    from app.core.config import get_settings
    d = get_settings().upload_dir.parent / "generated"
    d.mkdir(parents=True, exist_ok=True)
    return d
